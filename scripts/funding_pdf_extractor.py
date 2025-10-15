# scripts/funding_pdf_extractor.py
"""
Módulo principal de extracción de oportunidades de financiamiento
VERSIÓN FINAL CON EXTRACCIÓN HÍBRIDA (REGEX + GPT)
"""

import os
import re
import json
import time
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Set, Tuple
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openai import OpenAI
import sys
sys.path.append(str(Path(__file__).parent))

def get_config():
    """Obtiene la configuración actualizada"""
    import config
    import importlib
    importlib.reload(config)
    return config

def get_openai_client():
    """Obtiene cliente OpenAI con config actualizada"""
    cfg = get_config()
    return OpenAI(api_key=cfg.OPENAI_API_KEY)

# PROMPT ULTRA ESPECÍFICO
OPP_SYSTEM_PROMPT = """Eres experto en extraer información de Procurement Notices, especialmente de UNDP.

MISIÓN: Extraer TODO, especialmente DEADLINE y CONTACT que son campos CRÍTICOS.

CAMPOS OBLIGATORIOS (NUNCA en null):
1. **deadline**: BUSCA "DEADLINE", fechas como "17-Oct-25", "@", "AM", "PM"
2. **contact**: BUSCA emails con @ (especialmente @undp.org)
3. **sponsor**: Si ves "UNDP" → sponsor: "UNDP"
4. **amount**: Si no hay monto específico → "A determinar" (NO null)
5. **region**: Infiere de país mencionado
6. **country**: País específico si aparece

ESTRATEGIA AGRESIVA:
- "17-Oct-25 @ 01:59 AM" → deadline: "2025-10-17"
- "adquisiciones.sv@undp.org" → contact: "adquisiciones.sv@undp.org"
- "UNDP-SLV" → country: "El Salvador", sponsor: "UNDP"
- "EL SALVADOR" → country: "El Salvador", region: "América Latina"
- Sin monto → amount: "A determinar"
- Sin deadline claro → deadline: "unknown"

FORMATO JSON:
{
  "opportunities": [{
    "title": "string (OBLIGATORIO)",
    "summary": "string (OBLIGATORIO)",
    "sponsor": "UNDP u otra org",
    "amount": "Variable o A determinar si no hay monto",
    "currency": "USD por defecto",
    "deadline": "YYYY-MM-DD o fecha en formato ISO",
    "region": "América Latina, Global, etc.",
    "country": "País específico",
    "eligibility": "Quiénes pueden aplicar",
    "link": "URL si existe",
    "contact": "Email encontrado",
    "status": "open si deadline futuro, unknown si no está claro",
    "source_file": "nombre.pdf",
    "notes": "Info adicional relevante"
  }]
}

IMPORTANTE: Si la información preliminar ya provee deadline o contact, ÚSALA obligatoriamente."""

def read_pdf_text_enhanced(filepath: Path) -> str:
    """Extracción mejorada con LAParams para mejor detección"""
    try:
        # Método 1: Con LAParams optimizado
        laparams = LAParams(
            line_overlap=0.5,
            char_margin=2.0,
            line_margin=0.5,
            word_margin=0.1,
            boxes_flow=0.5,
            detect_vertical=False,
            all_texts=False
        )
        
        text = pdf_extract_text(str(filepath), laparams=laparams)
        
        if text and len(text) > 50:
            text = text.replace('\x0c', '\n')
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r' {2,}', ' ', text)
            return text.strip()
        
        # Método 2: Sin LAParams como fallback
        print("   ⚠️ Reintentando extracción básica...")
        text = pdf_extract_text(str(filepath))
        return text.strip() if text else ""
        
    except Exception as e:
        print(f"   ⚠️ Error en extracción: {e}")
        return ""

def extract_deadline_aggressive(text: str) -> Optional[str]:
    """Extracción agresiva de deadline - optimizada para UNDP"""
    
    # Patrón 1: "17-Oct-25 @ 01:59 AM" (formato UNDP típico)
    pattern1 = r'(\d{1,2})-([A-Z][a-z]{2})-(\d{2})\s*@\s*(\d{1,2}):(\d{2})\s*(AM|PM)'
    match = re.search(pattern1, text, re.IGNORECASE)
    if match:
        day, month_str, year, hour, minute, ampm = match.groups()
        
        months = {
            'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
            'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
            'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
        }
        month = months.get(month_str.lower()[:3], '01')
        full_year = f"20{year}"
        deadline = f"{full_year}-{month}-{day.zfill(2)}"
        return deadline
    
    # Patrón 2: Buscar "DEADLINE" y fecha cercana
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if 'deadline' in line.lower():
            context = '\n'.join(lines[i:min(i+3, len(lines))])
            
            date_patterns = [
                r'(\d{1,2})-([A-Z][a-z]{2})-(\d{2,4})',
                r'(\d{1,2})/(\d{1,2})/(\d{2,4})',
                r'(\d{4})-(\d{2})-(\d{2})'
            ]
            
            for pattern in date_patterns:
                match = re.search(pattern, context)
                if match:
                    parts = match.groups()
                    if len(parts) == 3:
                        try:
                            if '-' in match.group(0) and match.group(0)[2].isalpha():
                                day, month_str, year = parts
                                months = {
                                    'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
                                    'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
                                    'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
                                }
                                month = months.get(month_str.lower()[:3], '01')
                                full_year = f"20{year}" if len(year) == 2 else year
                                return f"{full_year}-{month}-{day.zfill(2)}"
                        except:
                            pass
    
    return None

def extract_contact_aggressive(text: str) -> Optional[str]:
    """Extracción agresiva de email de contacto"""
    
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if emails:
        # Priorizar emails de UNDP
        for email in emails:
            if 'undp' in email.lower():
                return email
        return emails[0]
    
    return None

def extract_reference_number(text: str) -> Optional[str]:
    """Extrae número de referencia formato UNDP"""
    pattern = r'(UNDP-[A-Z]{3}-\d{5})'
    match = re.search(pattern, text, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    return None

def clean_and_structure_text(text: str) -> str:
    """Limpia y marca secciones importantes del texto"""
    if not text:
        return ""
    
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Marcar secciones críticas con emojis para que GPT las identifique mejor
    text = re.sub(r'(DEADLINE[\s:]*)', r'\n\n⏰ DEADLINE CRÍTICO: ', text, flags=re.IGNORECASE)
    text = re.sub(r'(CONTACT[\s:]*)', r'\n\n📧 CONTACTO CRÍTICO: ', text, flags=re.IGNORECASE)
    text = re.sub(r'(REFERENCE\s+NUMBER[\s:]*)', r'\n\n🔢 REFERENCIA: ', text, flags=re.IGNORECASE)
    text = re.sub(r'(\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b)', r'\n📧 EMAIL: \1', text)
    
    return text

def extract_structured_info(text: str) -> Dict[str, str]:
    """Extracción estructurada con regex (complementa a GPT)"""
    info = {}
    
    # Deadline
    deadline = extract_deadline_aggressive(text)
    if deadline:
        info['deadline'] = deadline
        print(f"      🎯 Deadline regex: {deadline}")
    
    # Contact
    contact = extract_contact_aggressive(text)
    if contact:
        info['contact'] = contact
        print(f"      🎯 Contact regex: {contact}")
    
    # Reference
    ref = extract_reference_number(text)
    if ref:
        info['reference'] = ref
        print(f"      🎯 Referencia: {ref}")
    
    # Sponsor
    if 'UNDP' in text or 'undp' in text.lower():
        info['sponsor'] = 'UNDP'
    
    # País y región
    text_upper = text.upper()
    if 'EL SALVADOR' in text_upper or 'UNDP-SLV' in text:
        info['country'] = 'El Salvador'
        info['region'] = 'América Latina'
    elif 'GUATEMALA' in text_upper or 'UNDP-GTM' in text:
        info['country'] = 'Guatemala'
        info['region'] = 'América Latina'
    elif 'HONDURAS' in text_upper or 'UNDP-HND' in text:
        info['country'] = 'Honduras'
        info['region'] = 'América Latina'
    
    # URLs
    urls = re.findall(r'https?://[^\s<>"{}|\\^`\[\]]+', text)
    if urls:
        info['link'] = urls[0]
    
    return info

def keyword_focus(text: str, keywords: List[str]) -> str:
    """Prioriza párrafos con palabras clave"""
    if not text:
        return ""
    
    paragraphs = re.split(r'\n{2,}', text)
    relevant = []
    other = []
    
    for para in paragraphs:
        if any(kw.lower() in para.lower() for kw in keywords):
            relevant.append(para)
        else:
            other.append(para)
    
    if relevant:
        focused = '\n\n'.join(relevant)
        cfg = get_config()
        if len(focused) < cfg.CHUNK_SIZE * 3:
            focused += '\n\n' + '\n\n'.join(other[:5])
        return focused
    
    return text

def chunk_text(text: str, chunk_size: int, overlap: int, max_chunks: int) -> List[str]:
    """Divide texto en chunks con overlap"""
    if not text:
        return []
    
    words = text.split()
    if len(words) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(' '.join(words[start:end]))
        if end >= len(words):
            break
        start = end - overlap
    
    return chunks[:max_chunks]

def call_summary(text: str, filename: str, client: OpenAI, cfg) -> str:
    """Genera resumen ejecutivo del documento"""
    if not text:
        return "No se pudo extraer texto del documento."
    
    words = text.split()[:2000]
    text_limited = ' '.join(words)
    
    prompt = f"""Resume este documento en 120-180 palabras en {cfg.LANGUAGE_OUTPUT}.
Destaca: tema principal, propósito, y si contiene oportunidades de financiamiento.

Archivo: {filename}

{text_limited}"""
    
    try:
        response = client.chat.completions.create(
            model=cfg.OPENAI_MODEL,
            temperature=0.3,
            messages=[
                {"role": "system", "content": "Eres un experto en análisis de documentos."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generando resumen: {str(e)}"

def call_json_extract(text_chunk: str, filename: str, structured_info: Dict, client: OpenAI, cfg) -> Dict:
    """Extrae oportunidades con contexto de info ya encontrada"""
    if not text_chunk:
        return {"opportunities": []}
    
    # Crear hints con información ya extraída
    hints = "\n".join([f"- {k}: {v}" for k, v in structured_info.items()])
    
    user_prompt = f"""Archivo: {filename}
Idioma: {cfg.LANGUAGE_OUTPUT}

⚠️ INFORMACIÓN YA EXTRAÍDA (USA ESTO OBLIGATORIAMENTE):
{hints if hints else "No hay info preliminar"}

INSTRUCCIONES CRÍTICAS:
1. Si la info de arriba tiene deadline → ÚSALO (NO busques otro)
2. Si la info de arriba tiene contact → ÚSALO (NO busques otro)
3. Si la info de arriba tiene sponsor → ÚSALO
4. Para campos faltantes, extrae del texto
5. Si NO encuentras monto → amount: "A determinar"
6. Si NO encuentras currency → currency: "USD"

BUSCA ACTIVAMENTE:
- **title**: Título de la convocatoria
- **summary**: Resumen de 2-4 líneas
- **eligibility**: Quiénes pueden aplicar
- **notes**: Info adicional importante

TEXTO A ANALIZAR:
{text_chunk}

Devuelve JSON con opportunities. USA la info preliminar obligatoriamente."""
    
    for attempt in range(cfg.MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model=cfg.OPENAI_MODEL,
                temperature=cfg.OPENAI_TEMPERATURE,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": OPP_SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt}
                ]
            )
            
            result = json.loads(response.choices[0].message.content)
            
            if "opportunities" in result and isinstance(result["opportunities"], list):
                return result
            return {"opportunities": []}
                
        except json.JSONDecodeError:
            if attempt == cfg.MAX_RETRIES - 1:
                print(f"   ⚠️ Error parseando JSON")
                return {"opportunities": []}
        except Exception as e:
            if attempt == cfg.MAX_RETRIES - 1:
                print(f"   ⚠️ Error API: {str(e)}")
                return {"opportunities": []}
        
        time.sleep(cfg.RATE_LIMIT_DELAY)
    
    return {"opportunities": []}

def dedupe_opportunities(items: List[Dict]) -> List[Dict]:
    """Elimina oportunidades duplicadas"""
    seen = set()
    deduped = []
    
    for item in items:
        title = item.get("title", "").strip().lower()
        deadline = str(item.get("deadline", "")).strip()
        key = f"{title}|{deadline}"
        
        if key not in seen and title:
            seen.add(key)
            deduped.append(item)
    
    return deduped

def extract_opportunities_from_text(text: str, filename: str) -> Tuple[List[Dict], str]:
    """Pipeline completo de extracción"""
    if not text:
        return [], "Documento vacío o sin texto extraíble."
    
    cfg = get_config()
    client = get_openai_client()
    
    print(f"   📝 Texto extraído: {len(text)} caracteres")
    
    # Mostrar snippet para debug
    print(f"   📄 Primeros 150 caracteres:")
    print(f"      {text[:150].replace(chr(10), ' ')}")
    
    # Limpieza y estructuración
    print(f"   🧹 Limpiando y estructurando...")
    text = clean_and_structure_text(text)
    
    # REGEX PRIMERO (esto es clave)
    print(f"   🔍 Extrayendo con regex...")
    structured_info = extract_structured_info(text)
    if structured_info:
        print(f"   ✅ Regex encontró: {list(structured_info.keys())}")
    else:
        print(f"   ⚠️ Regex no encontró información clave")
    
    # Resumen
    print(f"   🤖 Generando resumen...")
    summary = call_summary(text, filename, client, cfg)
    
    # Filtro de keywords
    print(f"   🎯 Aplicando filtro de keywords...")
    focused_text = keyword_focus(text, cfg.KEYWORDS)
    
    # Chunking
    chunks = chunk_text(focused_text, cfg.CHUNK_SIZE, cfg.CHUNK_OVERLAP, cfg.MAX_CHUNKS_PER_DOC)
    print(f"   📦 Texto dividido en {len(chunks)} bloques")
    
    all_opportunities = []
    
    for i, chunk in enumerate(chunks, 1):
        print(f"   🔄 Analizando bloque {i}/{len(chunks)}...")
        
        # Pasar structured_info a GPT
        result = call_json_extract(chunk, filename, structured_info, client, cfg)
        opportunities = result.get("opportunities", [])
        
        # FORZAR campos críticos de regex
        for opp in opportunities:
            opp["source_file"] = filename
            
            # Forzar deadline y contact si regex los encontró
            for key in ['deadline', 'contact', 'sponsor', 'country', 'region', 'reference', 'link']:
                if key in structured_info:
                    if not opp.get(key) or opp.get(key) in [None, "null", "", "unknown"]:
                        opp[key] = structured_info[key]
                        print(f"      ✨ '{key}' FORZADO desde regex: {str(structured_info[key])[:50]}")
            
            # Valores por defecto para campos críticos
            if not opp.get('deadline') or opp.get('deadline') in [None, "null", ""]:
                opp['deadline'] = "unknown"
            
            if not opp.get('amount') or opp.get('amount') in [None, "null", ""]:
                opp['amount'] = "A determinar"
            
            if not opp.get('currency') or opp.get('currency') in [None, "null", ""]:
                opp['currency'] = "USD"
        
        all_opportunities.extend(opportunities)
        
        if i < len(chunks):
            time.sleep(cfg.RATE_LIMIT_DELAY)
    
    all_opportunities = dedupe_opportunities(all_opportunities)
    
    if not cfg.KEEP_CLOSED:
        all_opportunities = [
            opp for opp in all_opportunities
            if opp.get("status", "unknown").lower() != "closed"
        ]
    
    print(f"   ✅ {len(all_opportunities)} oportunidades encontradas")
    
    # Estadísticas de completitud
    for opp in all_opportunities:
        filled = sum(1 for v in opp.values() if v and v not in [None, "null", ""])
        total = len(opp)
        print(f"      📊 {opp.get('title', 'Sin título')[:35]}: {filled}/{total} campos ({int(filled/total*100)}%)")
        print(f"         💰 Amount: {opp.get('amount', 'N/A')}")
        print(f"         📅 Deadline: {opp.get('deadline', 'N/A')}")
        print(f"         📧 Contact: {opp.get('contact', 'N/A')}")
    
    return all_opportunities, summary

def process_pdf_folder(input_folder: Path = None, output_folder: Path = None) -> Dict:
    """Procesa todos los PDFs en una carpeta"""
    cfg = get_config()
    
    if input_folder is None:
        input_folder = cfg.PDFS_SALIDA
    if output_folder is None:
        output_folder = cfg.RESULTADOS
    
    print(f"\n📁 Carpeta de PDFs: {input_folder}")
    print(f"📁 Carpeta de resultados: {output_folder}")
    
    output_folder.mkdir(parents=True, exist_ok=True)
    
    pdf_files = list(input_folder.glob("*.pdf"))
    
    if not pdf_files:
        print("❌ No se encontraron PDFs en la carpeta")
        return {"error": "No PDFs found"}
    
    print(f"\n{'='*70}")
    print(f"📚 PROCESANDO {len(pdf_files)} PDFs")
    print(f"{'='*70}")
    
    all_results = []
    all_opportunities = []
    
    for idx, pdf_path in enumerate(pdf_files, 1):
        print(f"\n📄 [{idx}/{len(pdf_files)}] {pdf_path.name}")
        print(f"   {'-'*60}")
        
        text = read_pdf_text_enhanced(pdf_path)
        
        if not text or len(text) < 50:
            print(f"   ⚠️ No se pudo extraer texto suficiente")
            all_results.append({
                "filename": pdf_path.name,
                "summary": "No se pudo extraer texto del PDF",
                "opportunities_count": 0,
                "opportunities": []
            })
            continue
        
        opportunities, summary = extract_opportunities_from_text(text, pdf_path.name)
        
        result = {
            "filename": pdf_path.name,
            "summary": summary,
            "opportunities_count": len(opportunities),
            "opportunities": opportunities
        }
        
        all_results.append(result)
        all_opportunities.extend(opportunities)
        
        print(f"\n   📋 RESUMEN:")
        for line in summary.split('\n')[:3]:
            print(f"      {line}")
        
        if opportunities:
            print(f"\n   💰 {len(opportunities)} OPORTUNIDADES ENCONTRADAS:")
            for i, opp in enumerate(opportunities[:2], 1):
                print(f"      {i}. {opp.get('title', 'Sin título')[:60]}")
    
    # Guardar JSON
    json_output = {
        "processing_date": datetime.now().isoformat(),
        "total_pdfs": len(pdf_files),
        "total_opportunities": len(all_opportunities),
        "language": cfg.LANGUAGE_OUTPUT,
        "keep_closed": cfg.KEEP_CLOSED,
        "results": all_results
    }
    
    json_path = output_folder / "oportunidades_resultados.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_output, f, ensure_ascii=False, indent=2)
    
    # Crear DOCX
    docx_path = create_opportunities_docx(all_results, all_opportunities, output_folder)
    
    print(f"\n{'='*70}")
    print(f"✅ PROCESO COMPLETADO")
    print(f"{'='*70}")
    print(f"   • PDFs procesados: {len(pdf_files)}")
    print(f"   • Oportunidades encontradas: {len(all_opportunities)}")
    print(f"   • Archivo JSON: {json_path}")
    print(f"   • Documento Word: {docx_path}")
    
    return json_output

def create_opportunities_docx(results: List[Dict], all_opportunities: List[Dict], output_folder: Path) -> Path:
    """Crea documento Word con los resultados"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    title = doc.add_heading('REPORTE DE OPORTUNIDADES DE FINANCIAMIENTO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    cfg = get_config()
    doc.add_paragraph(f'Idioma: {"Español" if cfg.LANGUAGE_OUTPUT == "ES" else "English"}')
    doc.add_paragraph()
    
    doc.add_heading('RESUMEN EJECUTIVO', 1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'Valor'
    
    metrics = [
        ('Documentos analizados', str(len(results))),
        ('Oportunidades identificadas', str(len(all_opportunities))),
        ('Oportunidades abiertas', str(sum(1 for o in all_opportunities if o.get('status') == 'open'))),
        ('Oportunidades con deadline', str(sum(1 for o in all_opportunities if o.get('deadline') and o.get('deadline') != 'unknown')))
    ]
    
    for metric, value in metrics:
        row = table.add_row()
        row.cells[0].text = metric
        row.cells[1].text = value
    
    doc.add_page_break()
    
    if all_opportunities:
        doc.add_heading('TODAS LAS OPORTUNIDADES', 1)
        
        for i, opp in enumerate(all_opportunities, 1):
            p = doc.add_paragraph()
            runner = p.add_run(f"{i}. {opp.get('title', 'Sin título')}")
            runner.bold = True
            runner.font.size = Pt(12)
            
            detail_table = doc.add_table(rows=0, cols=2)
            detail_table.style = 'Table Grid'
            
            fields = [
                ('Resumen', opp.get('summary')),
                ('Patrocinador', opp.get('sponsor')),
                ('Monto', f"{opp.get('amount', '')} {opp.get('currency', '')}".strip() if opp.get('amount') else None),
                ('Fecha límite', opp.get('deadline')),
                ('Región', opp.get('region')),
                ('País', opp.get('country')),
                ('Elegibilidad', opp.get('eligibility')),
                ('Enlace', opp.get('link')),
                ('Contacto', opp.get('contact')),
                ('Estado', opp.get('status')),
                ('Archivo fuente', opp.get('source_file')),
                ('Notas', opp.get('notes'))
            ]
            
            for label, value in fields:
                if value and str(value).strip() and str(value) not in ["null", "None"]:
                    row = detail_table.add_row()
                    row.cells[0].text = label
                    row.cells[0].paragraphs[0].runs[0].bold = True
                    row.cells[0].width = Inches(2)
                    row.cells[1].text = str(value)
                    row.cells[1].width = Inches(4.5)
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    doc.add_heading('ANÁLISIS POR DOCUMENTO', 1)
    
    for result in results:
        doc.add_heading(f"📄 {result['filename']}", 2)
        
        doc.add_heading('Resumen del contenido:', 3)
        doc.add_paragraph(result['summary'])
        
        if result['opportunities_count'] > 0:
            doc.add_heading(f"Oportunidades encontradas ({result['opportunities_count']})", 3)
            
            for opp in result['opportunities']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(opp.get('title', 'Sin título')).bold = True
                if opp.get('deadline') and opp.get('deadline') != 'unknown':
                    p.add_run(f" - Deadline: {opp['deadline']}")
        else:
            doc.add_paragraph("No se encontraron oportunidades en este documento.", style='Intense Quote')
        
        doc.add_paragraph()
    
    docx_path = output_folder / "resumen_oportunidades.docx"
    doc.save(str(docx_path))
    
    return docx_path

if __name__ == "__main__":
    process_pdf_folder()